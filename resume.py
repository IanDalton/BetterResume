import os
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, Cm  # Add this import for font size and setting margins
from docx import Document
from bot import Bot
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH  # Add this import for justification


# This is only needed if you're using the builtin style above
def get_or_create_hyperlink_style(d):
    """If this document had no hyperlinks so far, the builtin
       Hyperlink style will likely be missing and we need to add it.
       There's no predefined value, different Word versions
       define it differently.
       This version is how Word 2019 defines it in the
       default theme, excluding a theme reference.
    """
    if "Hyperlink" not in d.styles:
        if "Default Character Font" not in d.styles:
            ds = d.styles.add_style("Default Character Font",
                                    docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                    True)
            ds.element.set(docx.oxml.shared.qn('w:default'), "1")
            ds.priority = 1
            ds.hidden = True
            ds.unhide_when_used = True
            del ds
        hs = d.styles.add_style("Hyperlink",
                                docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                True)
        hs.base_style = d.styles["Default Character Font"]
        hs.unhide_when_used = True
        hs.font.color.rgb = docx.shared.RGBColor(0x05, 0x63, 0xC1)
        hs.font.underline = True
        del hs

    return "Hyperlink"


def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.text.run.Run(
        docx.oxml.shared.OxmlElement('w:r'), paragraph)
    new_run.text = text

    # Set the run's style to the builtin hyperlink style, defining it if necessary
    new_run.style = get_or_create_hyperlink_style(part.document)
    # Alternatively, set the run's formatting explicitly
    # new_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    # new_run.font.underline = True

    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    return hyperlink


def set_paragraph_font(paragraph, font_name="Calibri", font_size=10):
    """Set font name and size for a paragraph."""
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)


def set_heading_font(paragraph, font_name="Times New Roman", font_size=16):
    """Set font name and size for a heading."""
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)


def set_paragraph_format(paragraph, line_spacing=1.0, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Set line spacing and alignment for a paragraph."""
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.alignment = alignment


class Resume():
    def __init__(self, JOB_PROMPT: str, csv_location: str = "jobs.csv", resume_name: str = None):
        self.bot = Bot()
        self.JOB_PROMPT = JOB_PROMPT
        data = pd.read_csv(csv_location)
        data["start_date"] = pd.to_datetime(
            data["start_date"], format="%d/%m/%Y")
        data["end_date"] = pd.to_datetime(data["end_date"], format="%d/%m/%Y")
        self.data = data
        self.response = None
        self.resume_name = resume_name if resume_name else f"resume_{"_".join(data[data["company"] == "name"]["description"].to_list()[0].split())}.pdf"
        if not self.resume_name.endswith(".pdf"):
            self.resume_name += ".pdf"

    def generate_resume(self, new_job_prompt: str = None):
        if not self.response or new_job_prompt:
            self.JOB_PROMPT = new_job_prompt if new_job_prompt else self.JOB_PROMPT
            self.response = self.bot.generate_response(self.JOB_PROMPT)
            if self.response["language"].upper() != "EN":
                self.response = self.bot.translate_response(self.response, self.JOB_PROMPT)
        

        data = self.data

        # Create the Word document
        
        document = Document()

        # Set reduced margins for the document
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(2)  # Set top margin to 1 cm
            section.bottom_margin = Cm(1)  # Set bottom margin to 1 cm
            section.left_margin = Cm(2)  # Set left margin to 1 cm
            section.right_margin = Cm(2)  # Set right margin to 1 cm

        heading = document.add_heading(
            f"{data[data['company'] == 'name']['description'].to_list()[0]} - {self.response['resume_section']['title']}", 0)
        set_heading_font(heading, font_name="Times New Roman", font_size=16)

        contact_paragraph = document.add_paragraph(
            f"{data[data['company'] == 'address']['description'].to_list()[0]} • {data[data['company'] == 'phone']['description'].to_list()[0]}")
        set_paragraph_font(contact_paragraph)
        set_paragraph_format(contact_paragraph)

        p = document.add_paragraph(
            f"{data[data['company'] == 'email']['description'].to_list()[0]} • ")
        set_paragraph_font(p)
        set_paragraph_format(p)
        websites = data[data["company"] == "website"]["description"].to_list()

        for i, website in enumerate(websites):
            add_hyperlink(p, website, website)
            if i < len(websites) - 1:
                p.add_run(" • ")
        summary_paragraph = document.add_paragraph(self.response["resume_section"]["professional_summary"])
        set_paragraph_font(summary_paragraph)
        set_paragraph_format(summary_paragraph)

        skills_heading = document.add_heading('SKILLS', level=0)
        set_heading_font(skills_heading, font_name="Times New Roman", font_size=11)

        skill: dict
        for skill in self.response["resume_section"].get("skills", []):
            p = document.add_paragraph()
            p.style = "List Bullet"
            set_paragraph_font(p)
            set_paragraph_format(p)
            p.add_run(skill["name"]).bold = True
            p.add_run(f" - {skill['description']}")
        experience_heading = document.add_heading('EXPERIENCE', level=0)
        set_heading_font(experience_heading, font_name="Times New Roman", font_size=11)

        experience: dict
        for experience in self.response["resume_section"].get("experience", []):
            p = document.add_paragraph()
            set_paragraph_font(p)
            set_paragraph_format(p)
            p.add_run(experience["company"]).bold = True
            p.add_run(
                f" • {experience['location']} • {experience['position']}")

            # Add a tab stop for right alignment
            tab_stop = p.paragraph_format.tab_stops.add_tab_stop(
                Inches(6.5))  # Adjust width as needed
            tab_stop.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            p.add_run(
                f"\t({experience['start_date']} - {experience['end_date'] if experience['end_date'] else 'Present'})")
            p = document.add_paragraph()
            set_paragraph_font(p)
            set_paragraph_format(p)
            p.add_run(experience["description"])

        education_heading = document.add_heading('EDUCATION AND CERTIFICATIONS', level=0)
        set_heading_font(education_heading, font_name="Times New Roman", font_size=11)

        for education in data[data["type"] == "education"].iterrows():
            p = document.add_paragraph(
                f"{education[1]['company']}, {education[1]['location']} • {education[1]['description']}")
            set_paragraph_font(p)
            set_paragraph_format(p)
            tab_stop = p.paragraph_format.tab_stops.add_tab_stop(
                Inches(6.5))  # Adjust width as needed
            tab_stop.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            p.add_run(
                f"\t{education[1]['start_date'].strftime('%m/%Y')} - {education[1]['end_date'].strftime('%m/%Y') if education[1]['end_date'] else 'Present'}")

        # Save the Word document
        docx_path = 'demo5.docx'
        document.save(docx_path)
        


        # Convert the Word document to PDF
        try:
            import comtypes.client

            def convert_to_pdf(docx_path, pdf_path):
                import comtypes.client
                try:
                    word = comtypes.client.CreateObject('Word.Application')
                    word.Visible = False  # Run Word in the background
                    doc = word.Documents.Open(os.path.abspath(docx_path))
                    # 17 is the PDF format
                    doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
                    doc.Close()
                    word.Quit()
                    print(f"Document successfully saved as {pdf_path}")
                except Exception as e:
                    print(f"Error during PDF conversion: {e}")
                    if 'word' in locals():
                        word.Quit()
                

            pdf_path = self.resume_name
            convert_to_pdf(docx_path, pdf_path)
            print(f"Document successfully saved as {pdf_path}")

        except ImportError:
            print(
                "comtypes library is required to convert to PDF. Install it using 'pip install comtypes'")
