import os
import logging
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from typing import Dict, Any, Optional
import pandas as pd
from .base_writer import BaseWriter
from models.resume import ResumeOutputFormat
from utils.word_utils import set_paragraph_font, set_paragraph_format, set_heading_font, add_hyperlink

class WordResumeWriter(BaseWriter):
    """
    A class for generating resumes in Word document format and optionally converting them to PDF.
    Methods
    -------
    write(response: dict, output: str = None, to_pdf: bool = False) -> Union[str, Document]:
        Generates a Word document based on the provided response data and optionally converts it to a PDF.
    generate_file(response: dict, output: str = None) -> Union[str, Document]:
        Creates a Word document with formatted resume content based on the provided response data.
    to_pdf(output: str, src_path: str = None) -> str:
        Converts a Word document to a PDF file using the comtypes library.
    """ 
    def __init__(self, template: str = None, csv_location: str = "jobs.csv", profile_image_path: Optional[str] = None):
        super().__init__(template, csv_location, ".docx", profile_image_path=profile_image_path)
        self._logger = logging.getLogger("betterresume.writer")



    def write(self,response:dict, output: str = None, to_pdf:bool=False):
        file = self.generate_file(response, output.replace(".pdf", ".docx") if output else None)
        self._logger.info("DOCX generated: %s", file if isinstance(file, str) else output)
        if not to_pdf:
            return file
        file = self.to_pdf(output.replace(".docx", ".pdf"), file)
        self._logger.info("PDF generated: %s", file)
        return file

    def generate_file(self,response:ResumeOutputFormat, output: str = None):
        
        data = self.data
        

        # Create the Word document
        
        document = Document()

        # Set reduced margins for the document
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(2)  # Set top margin to 1 cm
            section.bottom_margin = Cm(1)  # Set bottom margin to 1 cm
            section.left_margin = Cm(2)  # Set left margin to 1 cm
            section.right_margin = Cm(2)  # Set right margin to 1 cm

        # Helper safe getters
        def _safe_first(series, default: str = "") -> str:
            try:
                if series is None:
                    return default
                lst = series.to_list() if hasattr(series, "to_list") else list(series)
                return lst[0] if lst else default
            except Exception:
                return default

        
        
        resume_section = response.resume_section
        
        profile_path = self.profile_image_path if getattr(self, "profile_image_path", None) else None
        if profile_path and not os.path.isfile(profile_path):
            self._logger.debug("Profile image path does not exist: %s", profile_path)
            profile_path = None

        header_container = document
        table_created = None
        if profile_path:
            try:
                table_created = document.add_table(rows=1, cols=2)
                table_created.alignment = WD_TABLE_ALIGNMENT.LEFT
                table_created.autofit = False
                img_cell = table_created.rows[0].cells[0]
                text_cell = table_created.rows[0].cells[1]
                img_cell.width = Inches(1.75)
                text_cell.width = Inches(6.0)
                img_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                text_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                img_paragraph = img_cell.paragraphs[0]
                img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_paragraph.add_run()
                run.add_picture(profile_path, width=Inches(2.35))
                header_container = text_cell
            except Exception as e:
                self._logger.warning("Failed to insert profile image: %s", e)
                header_container = document
                if table_created is not None:
                    try:
                        tbl = table_created._tbl
                        tbl.getparent().remove(tbl)
                    except Exception:
                        pass

        def _add_heading(container, text: str):
            if not text:
                return None
            try:
                if hasattr(container, "add_heading"):
                    heading = container.add_heading(text, 0)
                else:
                    heading = container.add_paragraph(text)
                set_heading_font(heading, font_name="Times New Roman", font_size=16)
                return heading
            except Exception as exc:
                self._logger.debug("Failed adding heading: %s", exc)
                return None

        def _add_paragraph(container, text: str):
            if not text:
                return None
            try:
                paragraph = container.add_paragraph(text)
                set_paragraph_font(paragraph)
                set_paragraph_format(paragraph)
                return paragraph
            except Exception as exc:
                self._logger.debug("Failed adding paragraph: %s", exc)
                return None

        # Heading (Name - Title)
        try:
            name_txt = _safe_first(data[data['company'] == 'name']['description'], "")
            title_txt = resume_section.title
            heading_txt_parts = [p for p in [name_txt, title_txt] if p]
            if heading_txt_parts:
                _add_heading(header_container, " - ".join(heading_txt_parts))
        except Exception as e:
            self._logger.debug("Skipping heading due to missing data: %s", e)

        # Address • Phone
        try:
            address_txt = _safe_first(data[data['company'] == 'address']['description'], "")
            phone_txt = _safe_first(data[data['company'] == 'phone']['description'], "")
            parts = [p for p in [address_txt, phone_txt] if p]
            if parts:
                _add_paragraph(header_container, " • ".join(parts))
        except Exception as e:
            self._logger.debug("Skipping address/phone due to missing data: %s", e)

        # Email • Websites
        try:
            email_txt = _safe_first(data[data['company'] == 'email']['description'], "")
            websites = []
            try:
                websites = data[data["company"] == "website"]["description"].to_list()
            except Exception:
                websites = []
            if email_txt or websites:
                p = _add_paragraph(header_container, f"{email_txt}" + (" • " if email_txt and websites else ""))
                if p:
                    for i, website in enumerate([w for w in websites if w]):
                        try:
                            add_hyperlink(p, website, website)
                        except Exception:
                            p.add_run(website)
                        if i < len([w for w in websites if w]) - 1:
                            p.add_run(" • ")
        except Exception as e:
            self._logger.debug("Skipping email/websites due to missing data: %s", e)

        if header_container is not document:
            document.add_paragraph("")

        # Professional summary
        try:
            summary_txt = resume_section.professional_summary
            if summary_txt:
                summary_paragraph = document.add_paragraph(summary_txt)
                set_paragraph_font(summary_paragraph)
                set_paragraph_format(summary_paragraph)
        except Exception as e:
            self._logger.debug("Skipping professional summary: %s", e)

        # Skills
        try:
            skills = resume_section.skills
            if skills:
                skills_heading = document.add_heading('SKILLS', level=0)
                set_heading_font(skills_heading, font_name="Times New Roman", font_size=11)

                for skill in skills:
                    try:
                        p = document.add_paragraph()
                        p.style = "List Bullet"
                        set_paragraph_font(p)
                        set_paragraph_format(p)
                        name = skill.name
                        desc = skill.description
                        if name:
                            p.add_run(name).bold = True
                        if desc:
                            if name:
                                p.add_run(f" - {desc}")
                            else:
                                p.add_run(desc)
                    except Exception:
                        continue
        except Exception as e:
            self._logger.debug("Skipping skills due to missing data: %s", e)

        # Experience
        try:
            experiences = resume_section.experience
            if experiences:
                experience_heading = document.add_heading('EXPERIENCE', level=0)
                set_heading_font(experience_heading, font_name="Times New Roman", font_size=11)

                for experience in experiences:
                    try:
                        company = experience.company
                        location = experience.location
                        position = experience.position
                        start_date = experience.start_date
                        end_date = experience.end_date

                        p = document.add_paragraph()
                        set_paragraph_font(p)
                        set_paragraph_format(p)
                        parts_left = []
                        if company:
                            run = p.add_run(company)
                            run.bold = True
                        if location:
                            parts_left.append(location)
                        if position:
                            parts_left.append(position)
                        if parts_left:
                            p.add_run(" • " + " • ".join(parts_left))

                        # Add a tab stop for right alignment
                        try:
                            tab_stop = p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))  # Adjust width as needed
                            tab_stop.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        except Exception:
                            pass

                        date_right = None
                        if start_date or end_date:
                            end_txt = end_date if end_date and end_date.lower() != 'present' else 'Present'
                            date_right = f"\t({start_date} - {end_txt})"
                            p.add_run(date_right)

                        desc = experience.description
                        if desc:
                            p2 = document.add_paragraph()
                            set_paragraph_font(p2)
                            set_paragraph_format(p2)
                            p2.add_run(desc)
                    except Exception:
                        continue
        except Exception as e:
            self._logger.debug("Skipping experience due to missing data: %s", e)

        # Education and certifications
        try:
            education_list = resume_section.education
            if education_list:
                education_heading = document.add_heading('EDUCATION AND CERTIFICATIONS', level=0)
                set_heading_font(education_heading, font_name="Times New Roman", font_size=11)

                for edu in education_list:
                    try:
                        institution = edu.institution
                        degree = edu.degree
                        dates = edu.dates
                        
                        left_parts = [p for p in [institution, degree] if p]
                        left_txt = ", ".join(left_parts) if left_parts else institution or degree
                        line_txt = left_txt if left_txt else ""
                        
                        p = document.add_paragraph(line_txt)
                        set_paragraph_font(p)
                        set_paragraph_format(p)
                        try:
                            tab_stop = p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
                            tab_stop.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        except Exception:
                            pass
                        
                        if dates:
                            p.add_run(f"\t{dates}")
                    except Exception:
                        continue
        except Exception as e:
            self._logger.debug("Skipping education due to missing data: %s", e)
        if output:
            # Save the Word document
            document.save(output)
            return output
        
        return document
    def to_pdf(self, output: str, src_path: str= None):
        if not output.endswith(".pdf"):
            return output
        # Prefer LibreOffice (Linux container)
        try:
            import subprocess, shutil, sys
            if shutil.which("soffice") and src_path:
                out_dir = os.path.dirname(os.path.abspath(output))
                subprocess.run([
                    "soffice", "--headless", "--convert-to", "pdf", "--outdir", out_dir, os.path.abspath(src_path)
                ], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                # LibreOffice names file with .pdf in same dir; ensure expected name exists
                generated = os.path.join(out_dir, os.path.splitext(os.path.basename(src_path))[0] + ".pdf")
                if os.path.isfile(generated) and generated != os.path.abspath(output):
                    try:
                        os.replace(generated, output)
                    except Exception:
                        pass
                return output
        except Exception:
            pass
        # Fallback to Windows comtypes if available
        try:
            if not src_path:
                return output
            import comtypes.client  # type: ignore
            word = comtypes.client.CreateObject('Word.Application')
            word.Visible = False
            d = word.Documents.Open(os.path.abspath(src_path))
            d.SaveAs(os.path.abspath(output), FileFormat=17)
            d.Close(); word.Quit()
            return output
        except Exception:
            return output
