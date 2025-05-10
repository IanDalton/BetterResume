import os
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any
import pandas as pd
from .base_writer import BaseWriter
from utils.word_utils import set_paragraph_font, set_paragraph_format, set_heading_font, add_hyperlink

class WordResumeWriter(BaseWriter):
    """
    Take structured resume JSON and produce .docx and optional PDF.
    """
    

    def write(self,response:dict, output: str = None, to_pdf:bool=False):

        file = self.generate_file(response, output.replace(".pdf", ".docx") if output else None)
        if not to_pdf:
            return file
        file = self.to_pdf(output.replace(".docx", ".pdf"), file)
        return file

    def generate_file(self,response:dict, output: str = None):
        
        data = self.data
        self.response = response

        # Create the Word document
        
        document = Document()

        # Set reduced margins for the document
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(2)  # Set top margin to 1 cm
            section.bottom_margin = Cm(1)  # Set bottom margin to 1 cm
            section.left_margin = Cm(2)  # Set left margin to 1 cm
            section.right_margin = Cm(2)  # Set right margin to 1 cm

        heading = document.add_heading(
            f"{data[data['company'] == 'name']['description'].to_list()[0]} - {self.response['resume_section']['title']}", 0)
        set_heading_font(heading, font_name="Times New Roman", font_size=16)

        contact_paragraph = document.add_paragraph(
            f"{data[data['company'] == 'address']['description'].to_list()[0]} • {data[data['company'] == 'phone']['description'].to_list()[0]}")
        set_paragraph_font(contact_paragraph)
        set_paragraph_format(contact_paragraph)

        p = document.add_paragraph(
            f"{data[data['company'] == 'email']['description'].to_list()[0]} • ")
        set_paragraph_font(p)
        set_paragraph_format(p)
        websites = data[data["company"] == "website"]["description"].to_list()

        for i, website in enumerate(websites):
            add_hyperlink(p, website, website)
            if i < len(websites) - 1:
                p.add_run(" • ")
        summary_paragraph = document.add_paragraph(self.response["resume_section"]["professional_summary"])
        set_paragraph_font(summary_paragraph)
        set_paragraph_format(summary_paragraph)

        skills_heading = document.add_heading('SKILLS', level=0)
        set_heading_font(skills_heading, font_name="Times New Roman", font_size=11)

        skill: dict
        for skill in self.response["resume_section"].get("skills", []):
            p = document.add_paragraph()
            p.style = "List Bullet"
            set_paragraph_font(p)
            set_paragraph_format(p)
            p.add_run(skill["name"]).bold = True
            p.add_run(f" - {skill['description']}")
        experience_heading = document.add_heading('EXPERIENCE', level=0)
        set_heading_font(experience_heading, font_name="Times New Roman", font_size=11)

        experience: dict
        for experience in self.response["resume_section"].get("experience", []):
            p = document.add_paragraph()
            set_paragraph_font(p)
            set_paragraph_format(p)
            p.add_run(experience["company"]).bold = True
            p.add_run(
                f" • {experience['location']} • {experience['position']}")

            # Add a tab stop for right alignment
            tab_stop = p.paragraph_format.tab_stops.add_tab_stop(
                Inches(6.5))  # Adjust width as needed
            tab_stop.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(
                f"\t({experience['start_date']} - {experience['end_date'] if experience['end_date'] else 'Present'})")
            p = document.add_paragraph()
            set_paragraph_font(p)
            set_paragraph_format(p)
            p.add_run(experience["description"])

        education_heading = document.add_heading('EDUCATION AND CERTIFICATIONS', level=0)
        set_heading_font(education_heading, font_name="Times New Roman", font_size=11)

        for education in data[data["type"] == "education"].iterrows():
            p = document.add_paragraph(
                f"{education[1]['company']}, {education[1]['location']} • {education[1]['description']}")
            set_paragraph_font(p)
            set_paragraph_format(p)
            tab_stop = p.paragraph_format.tab_stops.add_tab_stop(
                Inches(6.5))  # Adjust width as needed
            tab_stop.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(
                f"\t{education[1]['start_date'].strftime('%m/%Y')} - {education[1]['end_date'].strftime('%m/%Y') if education[1]['end_date'] else 'Present'}")
        if output:
            # Save the Word document
            document.save(output)
            return output
        
        return document
    def to_pdf(self, output: str, src_path: str= None):
        try:

            if not output.endswith(".pdf"):
                raise ValueError("Output file must have a .pdf extension")
            import comtypes.client
            word = comtypes.client.CreateObject('Word.Application')
            word.Visible = False
            d = word.Documents.Open(os.path.abspath(src_path))
            d.SaveAs(os.path.abspath(output), FileFormat=17)
            d.Close(); word.Quit()

            return output
        except ImportError:
            raise RuntimeError("comtypes required for PDF conversion")
